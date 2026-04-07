"""
generate_report.py  -  Final Project Report (DOCX + PDF)
Real benchmark data: NVIDIA GeForce RTX 4070 Laptop GPU, CUDA 12.4, April 2026
"""

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Page setup ────────────────────────────────────────────────────────────────
sec = doc.sections[0]
sec.page_width    = Inches(8.27)
sec.page_height   = Inches(11.69)
sec.left_margin   = Inches(1.25)
sec.right_margin  = Inches(1.25)
sec.top_margin    = Inches(1.00)
sec.bottom_margin = Inches(0.67)

def add_page_numbers(section):
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.clear()
    run = p.add_run()
    for tag, txt in [('begin', None), (None, 'PAGE'), ('end', None)]:
        if tag:
            e = OxmlElement('w:fldChar'); e.set(qn('w:fldCharType'), tag); run._r.append(e)
        else:
            e = OxmlElement('w:instrText'); e.text = txt; run._r.append(e)
    run.font.name = 'Times New Roman'; run.font.size = Pt(12)

add_page_numbers(sec)

# ── Helpers ───────────────────────────────────────────────────────────────────
def sfont(run, size=12, bold=False, italic=False):
    run.font.name = 'Times New Roman'; run.font.size = Pt(size)
    run.bold = bold; run.italic = italic

def para(text='', align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12,
         bold=False, italic=False, sb=0, sa=6, indent=0):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    p.paragraph_format.line_spacing = Pt(18)
    if indent: p.paragraph_format.left_indent = Inches(indent)
    if text:
        r = p.add_run(text); sfont(r, size=size, bold=bold, italic=italic)
    return p

def ch(num, title):
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f'CHAPTER {num}'); sfont(r, 15, bold=True)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(0); p2.paragraph_format.space_after = Pt(14)
    r2 = p2.add_run(title.upper()); sfont(r2, 15, bold=True)

def sec_h(num, title):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(5)
    r = p.add_run(f'{num}  {title.upper()}'); sfont(r, 12, bold=True)

def subsec_h(num, title):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f'{num}  {title}'); sfont(r, 12, bold=True)

def tbl(headers, rows, col_w=None, caption=None):
    if caption:
        pc = doc.add_paragraph()
        pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pc.paragraph_format.space_before = Pt(8); pc.paragraph_format.space_after = Pt(4)
        rc = pc.add_run(caption); sfont(rc, 11, bold=True)
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hr.cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        pp = cell.paragraphs[0]; pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = pp.add_run(h); sfont(rr, 10, bold=True)
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),'D0D0D0')
        cell._tc.get_or_add_tcPr().append(shd)
    for ri, row_data in enumerate(rows):
        row = t.rows[ri+1]
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            pp = cell.paragraphs[0]; pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            rr = pp.add_run(str(val)); sfont(rr, 10)
    if col_w:
        for i, w in enumerate(col_w):
            for row in t.rows: row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t

def code(lines):
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Inches(0.3)
        r = p.add_run(line)
        r.font.name = 'Courier New'; r.font.size = Pt(9)

def rule():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(20)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),'6')
    bot.set(qn('w:space'),'1'); bot.set(qn('w:color'),'000000')
    pBdr.append(bot); pPr.append(pBdr)

# ══════════════════════════════════════════════════════════════════════════════
#  TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════
para('MANIPAL INSTITUTE OF TECHNOLOGY', WD_ALIGN_PARAGRAPH.CENTER, 16, bold=True, sb=0, sa=4)
para('MANIPAL', WD_ALIGN_PARAGRAPH.CENTER, 13, bold=True, sb=0, sa=2)
para('(A constituent unit of MAHE, Manipal)', WD_ALIGN_PARAGRAPH.CENTER, 11, sb=0, sa=2)
rule()
para('Mini-Project Report', WD_ALIGN_PARAGRAPH.CENTER, 13, bold=True, sb=0, sa=6)
for _ in range(2): para('', sa=2)
para('Optimizing CUDA Image Convolution via', WD_ALIGN_PARAGRAPH.CENTER, 18, bold=True, sb=0, sa=4)
para('Shared Memory Tiling and Separable Kernels', WD_ALIGN_PARAGRAPH.CENTER, 18, bold=True, sb=0, sa=30)
for _ in range(2): para('', sa=2)
para('SUBMITTED BY', WD_ALIGN_PARAGRAPH.CENTER, 12, bold=True, sb=0, sa=10)

at = doc.add_table(rows=4, cols=2)
at.style = 'Table Grid'; at.alignment = WD_TABLE_ALIGNMENT.CENTER
for ri, (nm, rg) in enumerate([('STUDENT NAME','REG. NO.'),
                                 ('Vasishta Nandipati','230962095'),
                                 ('Rishit Mathur','220962091'),
                                 ('Krishiv Kolanu','230962023')]):
    bd = (ri == 0)
    for ci, v in enumerate([nm, rg]):
        c = at.rows[ri].cells[ci]; pp = c.paragraphs[0]
        pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sfont(pp.add_run(v), 12, bold=bd)
for row in at.rows:
    row.cells[0].width = Inches(2.5); row.cells[1].width = Inches(1.5)

for _ in range(3): para('', sa=2)
para('Under the Guidance of:', WD_ALIGN_PARAGRAPH.CENTER, 12, bold=True, sb=0, sa=4)
para('Vidya Kamath', WD_ALIGN_PARAGRAPH.CENTER, 14, bold=True, sb=0, sa=3)
para('Assistant Professor', WD_ALIGN_PARAGRAPH.CENTER, 12, sb=0, sa=3)
para('School of Computer Engineering', WD_ALIGN_PARAGRAPH.CENTER, 12, sb=0, sa=3)
para('Manipal Institute of Technology, Manipal, Karnataka \u2013 576104', WD_ALIGN_PARAGRAPH.CENTER, 12, sb=0, sa=3)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  ABSTRACT
# ══════════════════════════════════════════════════════════════════════════════
para('ABSTRACT', WD_ALIGN_PARAGRAPH.CENTER, 15, bold=True, sb=12, sa=12)
para(
    "Image convolution underpins a large class of image processing tasks ranging from classical "
    "spatial filtering to the convolutional layers of modern deep learning networks. "
    "Although GPU hardware offers significant parallelism, a straightforward CUDA implementation "
    "of 2-D convolution makes no effort to reuse data cached on-chip and ends up issuing the "
    "same global memory reads repeatedly, leaving most of the device bandwidth idle. "
    "This project addresses that inefficiency through two optimisations. "
    "The first loads a tile of input pixels, together with a border region of width equal to "
    "the kernel radius, into shared memory once per thread block; subsequent computation then "
    "reads entirely from on-chip storage. "
    "The second decomposes a separable 2-D kernel into a horizontal and a vertical 1-D pass, "
    "each tiled independently, reducing the arithmetic work per output pixel from O(K\u00b2) to O(2K). "
    "Both optimisations are implemented from scratch in CUDA C++ and benchmarked on an "
    "NVIDIA GeForce RTX 4070 Laptop GPU against a sequential CPU reference and a naive GPU baseline. "
    "Tests span four kernel radii (r \u2208 {1, 3, 7, 15}) and four image resolutions "
    "(256\u00d7256 to 2048\u00d72048). "
    "The tiled kernel achieves up to 1.74\u00d7 better throughput than the naive kernel for the same "
    "kernel size, while the separable kernel yields the largest end-to-end speedups\u2014reaching "
    "8495\u00d7 over the CPU reference at the widest kernel and largest image tested. "
    "All three GPU implementations were verified to produce numerically correct results within "
    "floating-point tolerance across every test configuration."
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  CHAPTER 1
# ══════════════════════════════════════════════════════════════════════════════
ch(1, 'Introduction')

sec_h('1.1', 'General Introduction')
para(
    "Convolution is defined over an image I of height H and width W and a filter kernel K of "
    "size (2r+1)\u00d7(2r+1) as the weighted sum of the neighbourhood of each pixel. "
    "The computational cost of a direct implementation is O(H\u00b7W\u00b7K\u00b2), which "
    "for a 31\u00d731 kernel applied to a 2048\u00d72048 image amounts to roughly "
    "8 billion multiply-accumulate operations. "
    "Common uses include Gaussian blur for noise suppression, Sobel operators for edge detection, "
    "and the stacked convolutional layers of convolutional neural networks."
)
para(
    "NVIDIA GPUs execute thousands of threads simultaneously and expose a hierarchy of memory "
    "spaces with very different speeds. Global (device) memory offers the largest capacity but "
    "incurs around 300 clock cycles per access. Shared memory sits on the same chip as the "
    "compute cores, has roughly 5-cycle latency, and is partitioned per thread block. "
    "A naive convolution kernel ignores this distinction: every output thread independently "
    "fetches its entire K\u00d7K input neighbourhood from slow global memory, reading the "
    "same pixel up to K\u00b2 times. The result is a kernel that is strongly limited by "
    "memory bandwidth rather than by arithmetic capacity."
)
para(
    "Two widely studied techniques address this problem. Shared-memory tiling arranges "
    "threads within a block to load a contiguous tile of input data\u2014including a "
    "border region wide enough to cover the kernel overlap\u2014into shared memory before "
    "any arithmetic is performed, so that the expensive global reads happen only once per tile. "
    "Separable kernel decomposition exploits the fact that certain important kernels (Gaussian, "
    "box, binomial) factor into the outer product of two 1-D vectors, enabling the single "
    "2-D pass to be replaced by two cheaper 1-D passes, each tiled in shared memory."
)
para(
    "This project implements all three approaches (naive, tiled, separable), verifies their "
    "correctness against a CPU reference, and measures their performance across a range of "
    "kernel sizes and image resolutions on an RTX 4070 Laptop GPU using CUDA 12.4."
)

sec_h('1.2', 'Organization')
para(
    "Chapter 2 states the problem and its computational context. Chapter 3 lists the specific "
    "objectives. Chapter 4 covers related prior work. Chapter 5 describes the methodology and "
    "benchmark design. Chapter 6 presents the full implementation and the measured results. "
    "Chapter 7 records individual contributions. References follow at the end."
)

sec_h('1.3', 'Area of Computer Science')
para(
    "The work spans parallel computing (GPU programming, SIMT execution, shared-memory "
    "synchronisation), computer architecture (memory hierarchy, roofline performance model, "
    "occupancy analysis), and computer vision (spatial image filtering). "
    "The primary topics from the Parallel Computing and Architecture course that this project "
    "exercises are shared memory, thread-block tiling, and the roofline model."
)

sec_h('1.4', 'Hardware and Software Requirements')

para('Hardware', bold=True, sa=4, sb=6)
tbl(['Component', 'Specification'],
    [('GPU', 'NVIDIA GeForce RTX 4070 Laptop GPU (Ada Lovelace, CC 8.9)'),
     ('GPU Memory', '8 GB GDDR6'),
     ('Peak Memory Bandwidth', '256 GB/s'),
     ('Shared Memory per Block', '48 KB (default)'),
     ('Streaming Multiprocessors', '36 SMs'),
     ('Host CPU', 'x86-64, 8+ cores'),
     ('Host RAM', '16 GB')],
    col_w=[1.8, 4.0])

para('Software', bold=True, sa=4, sb=6)
tbl(['Component', 'Version'],
    [('Operating System', 'Ubuntu 22.04 LTS'),
     ('CUDA Toolkit', '12.4'),
     ('NVCC Compiler', '12.4 (V12.4.131)'),
     ('NVIDIA Driver', '550.163.01'),
     ('GCC / G++', '12.x'),
     ('Build System', 'GNU Make 4.x')],
    col_w=[2.5, 3.0])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  CHAPTER 2
# ══════════════════════════════════════════════════════════════════════════════
ch(2, 'Problem Definition')
para(
    "A naive CUDA implementation of 2-D convolution assigns one output pixel to each GPU thread. "
    "Each thread independently computes the inner product of its K\u00d7K neighbourhood with the "
    "filter kernel, fetching every required input pixel directly from global memory. "
    "Consider a 16\u00d716 output tile processed by a single thread block. With a 15\u00d715 kernel "
    "(r = 7) each of the 256 threads requires 225 input values, for a total of 57,600 global "
    "memory reads per block. Yet the unique inputs needed cover only a 30\u00d730 = 900-element "
    "apron region\u2014meaning the same data is read an average of 64 times from slow DRAM."
)
para(
    "The degree of inefficiency is characterised by arithmetic intensity (AI), the ratio of "
    "floating-point operations to bytes of memory traffic. For the naive 2-D kernel with radius r:"
)
para("AI\u2082D  =  2 \u00b7 (2r+1)\u00b2  /  (2 \u00b7 4)  FLOP/byte",
     WD_ALIGN_PARAGRAPH.CENTER, italic=False, sb=4, sa=4)
para(
    "At r = 1 (K = 3) this gives 2.25 FLOP/byte, far below the ridge point of the RTX 4070 "
    "Laptop GPU (~32 FLOP/byte based on peak 16.1 TFLOP/s and 256 GB/s bandwidth). "
    "Measured results confirmed this: the naive kernel for r = 1 attained only 529 GFLOP/s on a "
    "512\u00d7512 image against a theoretical peak exceeding 16,000 GFLOP/s. "
    "Even for large kernels (r = 15, AI = 240 FLOP/byte) the naive kernel is inefficient because "
    "it brings the same data from global memory repeatedly instead of sharing it through on-chip storage."
)
para(
    "The project therefore investigates whether shared-memory tiling and separable decomposition "
    "can close this gap, by how much, and for which kernel-size regimes each optimisation is most effective."
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  CHAPTER 3
# ══════════════════════════════════════════════════════════════════════════════
ch(3, 'Objectives')
sec_h('3.1', 'Baseline Implementation')
para(
    "Implement a naive CUDA convolution kernel in which every thread independently reads its "
    "K\u00d7K neighbourhood from global memory. Verify correctness and establish baseline "
    "execution times, effective bandwidth, and GFLOP/s figures."
)
sec_h('3.2', 'Shared Memory Tiling')
para(
    "Design a tiled kernel that loads the K/2-pixel halo together with the output tile into "
    "shared memory once per block, eliminating redundant global reads within that tile. "
    "Confirm that the halo and boundary handling (border-replicate padding) is correct for "
    "all edge and corner pixels."
)
sec_h('3.3', 'Separable Kernel Decomposition')
para(
    "Implement a two-pass separable pipeline\u2014horizontal then vertical\u2014each pass "
    "independently tiled in shared memory. Use the decomposition to reduce per-pixel arithmetic "
    "from O(K\u00b2) to O(2K) and verify the result matches the 2-D reference."
)
sec_h('3.4', 'Correctness Validation')
para(
    "For every (kernel radius, image resolution) combination, compare each GPU output against "
    "the CPU reference and confirm that the maximum absolute difference is below 10\u207b\u00b3. "
    "Report per-run error to detect any border-handling defect."
)
sec_h('3.5', 'Performance Benchmarking')
para(
    "Benchmark all four methods (CPU, Naive GPU, Tiled GPU, Separable GPU) across radii "
    "r \u2208 {1, 3, 7, 15} and resolutions {256\u00d7256, 512\u00d7512, 1024\u00d71024, 2048\u00d72048}. "
    "Report execution time in ms, effective memory bandwidth in GB/s, arithmetic throughput "
    "in GFLOP/s, and speedup relative to the CPU reference."
)
sec_h('3.6', 'Roofline Analysis')
para(
    "Compute the arithmetic intensity of each method for each kernel radius and compare the "
    "achieved performance against the device\u2019s memory-bandwidth and compute rooflines, "
    "identifying whether each configuration is bandwidth-limited or compute-limited."
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  CHAPTER 4
# ══════════════════════════════════════════════════════════════════════════════
ch(4, 'Background')
para(
    "GPU-accelerated convolution has been studied since the early days of general-purpose GPU "
    "computing. Podlozhnyuk [1] published one of the first detailed treatments in 2007, showing "
    "that storing filter coefficients in constant memory reduces the effective access cost to "
    "near zero because all threads in a warp read the same coefficient at each step, triggering "
    "a broadcast from the constant cache. Kirk and Hwu [2] systematised the tiling strategy "
    "in the context of a broader parallel patterns framework and established that on-chip data "
    "reuse is the dominant factor in GPU memory-bandwidth efficiency."
)
para(
    "Separable filter decomposition has a long history in signal processing. Getreuer [3] "
    "surveys Gaussian filter algorithms and shows that a separable 2-D implementation "
    "consistently beats a direct 2-D implementation once the kernel radius exceeds roughly "
    "two pixels, because the arithmetic savings outweigh any additional overhead from managing "
    "an intermediate buffer. On the GPU the two passes can be overlapped using CUDA streams "
    "to reduce overall latency further."
)
para(
    "More recent hardware-specific work by Hong et al. [4] introduced warp-shuffle reductions "
    "to replace shared-memory reductions that are prone to bank conflicts, reporting up to "
    "1.8\u00d7 additional gain over standard tiling on Ampere-class devices. At the library "
    "level, NVIDIA cuDNN [6] selects among direct, FFT-based, Winograd, and implicit-GEMM "
    "algorithms at runtime. This project deliberately works without cuDNN to build direct "
    "experience with kernel-level optimisation and to treat cuDNN as an upper-bound reference "
    "rather than a dependency."
)
para(
    "The roofline model of Williams et al. [9] provides a concise visual framework that "
    "relates arithmetic intensity to achievable performance and identifies whether a given "
    "kernel is limited by memory bandwidth or by floating-point throughput. This model "
    "guides the analysis in Chapter 6."
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  CHAPTER 5
# ══════════════════════════════════════════════════════════════════════════════
ch(5, 'Methodology')

sec_h('5.1', 'Overall Approach')
para(
    "The project follows a build-and-compare methodology. A sequential CPU implementation "
    "serves as the numerical reference. Three CUDA kernels are written from scratch, each "
    "building on the previous: first the naive baseline, then the tiled variant, then the "
    "separable pipeline. A common benchmark harness times each method under identical "
    "conditions and compares the output against the CPU reference pixel by pixel."
)

sec_h('5.2', 'Tile and Block Dimensions')
para(
    "Thread blocks of 16\u00d716 = 256 threads are used throughout. This choice keeps the "
    "shared memory allocation well under the 48 KB per-block limit even at the maximum kernel "
    "radius of 15: the tiled 2-D kernel needs (16+30)\u00d7(16+30)\u00d74 = 8,464 bytes per "
    "block, which is 17.6% of the limit. The 256-thread block also supports reasonable occupancy "
    "on the RTX 4070 Laptop GPU, which allows up to 1,536 threads per SM."
)

sec_h('5.3', 'Kernel Coefficients in Constant Memory')
para(
    "Convolution filter weights are stored in CUDA constant memory (64 KB device-wide, "
    "cached and broadcast-optimised). At each step of the convolution inner loop all threads "
    "in a warp read the same coefficient, which is served from the constant cache as a single "
    "broadcast transaction rather than 32 separate global reads. The 2-D kernels occupy at "
    "most 31\u00d731\u00d74 = 3,844 bytes; the 1-D kernel for the separable pass occupies "
    "at most 31\u00d74 = 124 bytes."
)

sec_h('5.4', 'Border Handling')
para(
    "Out-of-bounds coordinates are clamped to the nearest valid pixel using "
    "min(max(index, 0), size\u22121) (border-replicate / BORDER_REPLICATE in OpenCV terminology). "
    "The clamping is applied during the shared-memory load phase for tiled kernels so that "
    "the arithmetic phase reads only from shared memory with no per-pixel branching."
)

sec_h('5.5', 'Gaussian Kernel Generation')
para(
    "For each tested radius r the 1-D Gaussian kernel h is computed as "
    "h[k] = exp(\u2212(k\u2212r)\u00b2 / (2\u03c3\u00b2)), with \u03c3 = r/2, "
    "and normalised to unit sum. The 2-D kernel is the outer product H[i][j] = h[i]\u00b7h[j], "
    "also normalised. The same h is used for both passes of the separable kernel."
)

sec_h('5.6', 'Benchmark Protocol')
tbl(['Parameter', 'Value'],
    [('Warmup launches before timing', '3'),
     ('Timed launches per measurement', '10'),
     ('Reported figure', 'Arithmetic mean of 10 runs'),
     ('GPU timer', 'cudaEvent_t (GPU-side, ~0.5 \u03bcs resolution)'),
     ('CPU timer', 'std::chrono::high_resolution_clock'),
     ('Correctness threshold', 'max absolute difference < 10\u207b\u00b3'),
     ('Effective BW formula', '2 \u00b7 H \u00b7 W \u00b7 4 / time_ms / 10\u2076  (GB/s)'),
     ('GFLOP/s (2-D kernel)', '2 \u00b7 H \u00b7 W \u00b7 K\u00b2 / time_ms / 10\u2076'),
     ('GFLOP/s (separable)',   '2 \u00b7 H \u00b7 W \u00b7 2K / time_ms / 10\u2076')],
    col_w=[2.8, 3.5])

sec_h('5.7', 'Test Input')
para(
    "A synthetic floating-point image is generated at runtime using a sum of sine and cosine "
    "waves at different spatial frequencies. Values are mapped to [0, 1]. Using a generated "
    "rather than a file-loaded image removes any dependency on external image libraries and "
    "ensures fully reproducible inputs across runs."
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  CHAPTER 6
# ══════════════════════════════════════════════════════════════════════════════
ch(6, 'Implementation and Results')

sec_h('6.1', 'Source File Structure')
tbl(['File', 'Contents'],
    [('include/convolution.h',    'Tile constants (TILE_W = TILE_H = 16), CUDA_CHECK macro, function prototypes'),
     ('src/naive_conv.cu',         'Global-memory baseline kernel and launch wrapper'),
     ('src/tiled_conv.cu',         'Shared-memory tiled kernel and launch wrapper'),
     ('src/separable_conv.cu',     'Horizontal + vertical 1-D kernels and launch wrapper'),
     ('src/main.cu',               'CPU reference, Gaussian generator, benchmark harness, main()')],
    col_w=[2.0, 4.3])

sec_h('6.2', 'Naive Convolution Kernel')
para(
    "Each thread computes one output pixel. The inner loop iterates over all K\u00d7K kernel "
    "positions, reads the corresponding input pixel from global memory (with clamping), "
    "multiplies by the constant-memory coefficient, and accumulates."
)
code([
    "__global__ void naiveConvKernel(",
    "    const float* __restrict__ input, float* __restrict__ output,",
    "    int width, int height, int kernelRadius)",
    "{",
    "    int col = blockIdx.x * blockDim.x + threadIdx.x;",
    "    int row = blockIdx.y * blockDim.y + threadIdx.y;",
    "    if (col >= width || row >= height) return;",
    "    const int kSize = 2 * kernelRadius + 1;",
    "    float sum = 0.0f;",
    "    for (int kr = 0; kr < kSize; ++kr)",
    "        for (int kc = 0; kc < kSize; ++kc) {",
    "            int r = min(max(row + kr - kernelRadius, 0), height - 1);",
    "            int c = min(max(col + kc - kernelRadius, 0), width  - 1);",
    "            sum += input[r * width + c] * d_kernelNaive[kr * kSize + kc];",
    "        }",
    "    output[row * width + col] = sum;",
    "}",
])
doc.add_paragraph()

sec_h('6.3', 'Tiled Shared-Memory Kernel')
para(
    "The tiled kernel adds a cooperative load phase before the computation. All 256 threads in "
    "a block together fill a (TILE_W + 2r) \u00d7 (TILE_H + 2r) shared-memory buffer, using a "
    "stride-based loop so every element of the apron is covered. A __syncthreads() barrier "
    "separates the load from the compute phase, ensuring the buffer is fully populated before "
    "any thread begins its convolution loop."
)
code([
    "// Cooperative apron load into shared memory",
    "for (int dy = ty; dy < smH; dy += TILE_H)",
    "    for (int dx = tx; dx < smW; dx += TILE_W) {",
    "        int r = min(max(baseRow + dy, 0), height - 1);",
    "        int c = min(max(baseCol + dx, 0), width  - 1);",
    "        smem[dy * smW + dx] = input[r * width + c];",
    "    }",
    "__syncthreads();",
    "",
    "// Convolution reads entirely from shared memory",
    "float sum = 0.0f;",
    "for (int kr = 0; kr < kSize; ++kr)",
    "    for (int kc = 0; kc < kSize; ++kc)",
    "        sum += smem[(ty + kr) * smW + (tx + kc)]",
    "               * d_kernelTiled[kr * kSize + kc];",
    "output[outRow * width + outCol] = sum;",
])
doc.add_paragraph()

sec_h('6.4', 'Separable Convolution Kernels')
para(
    "The separable pipeline runs two separate kernel launches. The horizontal kernel loads "
    "TILE_H rows, each padded with r elements on each side, into shared memory and convolves "
    "every row with the 1-D filter h, writing results to an intermediate device buffer d_temp. "
    "The vertical kernel then loads columns from d_temp\u2014each padded with r elements above "
    "and below\u2014and applies h along the column direction."
)
para(
    "The index arithmetic for the horizontal pass is: smem[ty * smW + tx + k] corresponds to "
    "input pixel at column (outCol \u2212 r + k) in row outRow, so the 1-D sum "
    "\u03a3 smem[ty*smW + tx + k] \u00b7 h[k] reproduces the correct row convolution. "
    "An identical argument applies to the vertical pass with rows replacing columns."
)
code([
    "// Horizontal pass  (input -> d_temp)",
    "const size_t hSmem = TILE_H * (TILE_W + 2*kernelRadius) * sizeof(float);",
    "sepHorizontalKernel<<<grid, block, hSmem>>>(input, d_temp, ...);",
    "cudaDeviceSynchronize();",
    "",
    "// Vertical pass    (d_temp -> output)",
    "const size_t vSmem = (TILE_H + 2*kernelRadius) * TILE_W * sizeof(float);",
    "sepVerticalKernel  <<<grid, block, vSmem>>>(d_temp, output, ...);",
])
doc.add_paragraph()

sec_h('6.5', 'Measured Results')

# ── Actual output from the benchmark ─────────────────────────────────────────
para(
    "The following results were obtained by running the compiled benchmark binary on the "
    "NVIDIA GeForce RTX 4070 Laptop GPU (CUDA 12.4, driver 550.163.01). "
    "All GPU implementations passed the correctness check (max error < 10\u207b\u00b3) "
    "for every tested configuration."
)

para("Table 6.1   Full Benchmark Results \u2014 All Resolutions and Kernel Radii",
     WD_ALIGN_PARAGRAPH.CENTER, 11, bold=True, sb=8, sa=4)

# All real data
tbl(
    ['Resolution', 'Method', 'r', 'Time (ms)', 'BW (GB/s)', 'GFLOP/s', 'CPU Speedup'],
    [
        # 256x256
        ('256\u00d7256','CPU Reference',  1, '0.673',  '\u2014',   '1.75',    '1.00\u00d7'),
        ('256\u00d7256','Naive GPU',       1, '0.005',  '104.5',  '235.1',  '134.0\u00d7'),
        ('256\u00d7256','Tiled GPU',       1, '0.006',   '91.4',  '205.7',  '117.3\u00d7'),
        ('256\u00d7256','Separable GPU',   1, '0.013',   '40.3',   '60.5',   '51.7\u00d7'),
        ('256\u00d7256','CPU Reference',   3, '2.678',  '\u2014',   '2.40',    '1.00\u00d7'),
        ('256\u00d7256','Naive GPU',       3, '0.010',   '50.7',  '621.0',  '258.9\u00d7'),
        ('256\u00d7256','Tiled GPU',       3, '0.008',   '62.4',  '764.9',  '318.9\u00d7'),
        ('256\u00d7256','Separable GPU',   3, '0.013',   '39.7',  '138.9',  '202.7\u00d7'),
        ('256\u00d7256','CPU Reference',   7,'11.189',  '\u2014',   '2.64',    '1.00\u00d7'),
        ('256\u00d7256','Naive GPU',       7, '0.030',   '17.4',  '976.4',  '370.4\u00d7'),
        ('256\u00d7256','Tiled GPU',       7, '0.020',   '26.0', '1461.9',  '554.7\u00d7'),
        ('256\u00d7256','Separable GPU',   7, '0.014',   '36.3',  '272.3',  '775.0\u00d7'),
        ('256\u00d7256','CPU Reference',  15,'48.948',  '\u2014',   '2.57',    '1.00\u00d7'),
        ('256\u00d7256','Naive GPU',      15, '0.113',    '4.6', '1111.2',  '431.8\u00d7'),
        ('256\u00d7256','Tiled GPU',      15, '0.072',    '7.3', '1759.8',  '683.8\u00d7'),
        ('256\u00d7256','Separable GPU',  15, '0.016',   '31.8',  '492.9', '2969.0\u00d7'),
        # 512x512
        ('512\u00d7512','CPU Reference',   1, '2.699',  '\u2014',   '1.75',    '1.00\u00d7'),
        ('512\u00d7512','Naive GPU',        1, '0.009',  '235.4',  '529.7',  '303.0\u00d7'),
        ('512\u00d7512','Tiled GPU',        1, '0.010',  '203.3',  '457.4',  '261.6\u00d7'),
        ('512\u00d7512','Separable GPU',    1, '0.018',  '114.5',  '171.7',  '147.3\u00d7'),
        ('512\u00d7512','CPU Reference',    3,'10.641',  '\u2014',   '2.41',    '1.00\u00d7'),
        ('512\u00d7512','Naive GPU',        3, '0.026',   '80.0',  '980.0',  '405.9\u00d7'),
        ('512\u00d7512','Tiled GPU',        3, '0.020',  '104.5', '1280.0',  '530.2\u00d7'),
        ('512\u00d7512','Separable GPU',    3, '0.020',  '106.2',  '371.6',  '538.7\u00d7'),
        ('512\u00d7512','CPU Reference',    7,'44.775',  '\u2014',   '2.63',    '1.00\u00d7'),
        ('512\u00d7512','Naive GPU',        7, '0.097',   '21.7', '1220.3',  '463.2\u00d7'),
        ('512\u00d7512','Tiled GPU',        7, '0.061',   '34.2', '1926.4',  '731.2\u00d7'),
        ('512\u00d7512','Separable GPU',    7, '0.023',   '91.0',  '682.7', '1943.4\u00d7'),
        ('512\u00d7512','CPU Reference',   15,'194.12',  '\u2014',   '2.60',    '1.00\u00d7'),
        ('512\u00d7512','Naive GPU',       15, '0.418',    '5.0', '1205.1',  '464.3\u00d7'),
        ('512\u00d7512','Tiled GPU',       15, '0.240',    '8.7', '2097.3',  '808.1\u00d7'),
        ('512\u00d7512','Separable GPU',   15, '0.052',   '40.2',  '622.4', '3717.1\u00d7'),
        # 1024x1024
        ('1024\u00d71024','CPU Reference',  1, '10.706',  '\u2014',   '1.76',    '1.00\u00d7'),
        ('1024\u00d71024','Naive GPU',       1,  '0.025',  '341.3',  '768.0',  '435.6\u00d7'),
        ('1024\u00d71024','Tiled GPU',       1,  '0.035',  '237.4',  '534.3',  '303.0\u00d7'),
        ('1024\u00d71024','Separable GPU',   1,  '0.042',  '198.8',  '298.3',  '253.8\u00d7'),
        ('1024\u00d71024','CPU Reference',   3, '42.527',  '\u2014',   '2.42',    '1.00\u00d7'),
        ('1024\u00d71024','Naive GPU',       3,  '0.091',   '92.1', '1128.8',  '467.2\u00d7'),
        ('1024\u00d71024','Tiled GPU',       3,  '0.067',  '125.5', '1536.8',  '636.0\u00d7'),
        ('1024\u00d71024','Separable GPU',   3,  '0.048',  '173.2',  '606.2',  '878.0\u00d7'),
        ('1024\u00d71024','CPU Reference',   7,'184.220',  '\u2014',   '2.56',    '1.00\u00d7'),
        ('1024\u00d71024','Naive GPU',       7,  '0.367',   '22.9', '1285.4',  '501.8\u00d7'),
        ('1024\u00d71024','Tiled GPU',       7,  '0.228',   '36.8', '2068.2',  '807.5\u00d7'),
        ('1024\u00d71024','Separable GPU',   7,  '0.063',  '133.6', '1002.3', '2934.8\u00d7'),
        ('1024\u00d71024','CPU Reference',  15,'779.465',  '\u2014',   '2.59',    '1.00\u00d7'),
        ('1024\u00d71024','Naive GPU',      15,  '1.495',    '5.6', '1348.1',  '521.4\u00d7'),
        ('1024\u00d71024','Tiled GPU',      15,  '0.927',    '9.1', '2174.3',  '840.9\u00d7'),
        ('1024\u00d71024','Separable GPU',  15,  '0.092',   '90.9', '1409.3', '8448.4\u00d7'),
        # 2048x2048
        ('2048\u00d72048','CPU Reference',  1, '43.268',  '\u2014',   '1.74',    '1.00\u00d7'),
        ('2048\u00d72048','Naive GPU',       1,  '0.090',  '374.5',  '842.6',  '482.9\u00d7'),
        ('2048\u00d72048','Tiled GPU',       1,  '0.116',  '290.2',  '653.0',  '374.3\u00d7'),
        ('2048\u00d72048','Separable GPU',   1,  '0.268',  '125.3',  '188.0',  '161.6\u00d7'),
        ('2048\u00d72048','CPU Reference',   3,'170.698',  '\u2014',   '2.41',    '1.00\u00d7'),
        ('2048\u00d72048','Naive GPU',       3,  '0.352',   '95.4', '1168.6',  '485.3\u00d7'),
        ('2048\u00d72048','Tiled GPU',       3,  '0.257',  '130.7', '1601.2',  '664.9\u00d7'),
        ('2048\u00d72048','Separable GPU',   3,  '0.317',  '105.8',  '370.3',  '538.3\u00d7'),
        ('2048\u00d72048','CPU Reference',   7,'737.687',  '\u2014',   '2.56',    '1.00\u00d7'),
        ('2048\u00d72048','Naive GPU',       7,  '1.450',   '23.1', '1301.3',  '508.6\u00d7'),
        ('2048\u00d72048','Tiled GPU',       7,  '0.900',   '37.3', '2096.5',  '819.4\u00d7'),
        ('2048\u00d72048','Separable GPU',   7,  '0.276',  '121.7',  '912.9', '2676.1\u00d7'),
        ('2048\u00d72048','CPU Reference',  15,'3101.35',  '\u2014',   '2.60',    '1.00\u00d7'),
        ('2048\u00d72048','Naive GPU',      15,  '6.095',    '5.5', '1322.6',  '508.8\u00d7'),
        ('2048\u00d72048','Tiled GPU',      15,  '3.682',    '9.1', '2189.3',  '842.2\u00d7'),
        ('2048\u00d72048','Separable GPU',  15,  '0.365',   '91.9', '1424.7', '8495.5\u00d7'),
    ],
    col_w=[1.15, 1.3, 0.3, 0.85, 0.85, 0.85, 1.05]
)

sec_h('6.6', 'Discussion of Results')
para(
    "Several consistent patterns emerge from the data."
)

subsec_h('6.6.1', 'Small Kernels (r = 1)')
para(
    "For the smallest kernel tested (3\u00d73, r = 1), the naive kernel is actually faster than "
    "the tiled kernel at all resolutions. This is because a 3\u00d73 neighbourhood has very "
    "little spatial overlap between adjacent output pixels: the apron adds only 2 border rows "
    "and 2 border columns to the 16\u00d716 tile, so the global memory traffic reduction is "
    "modest while the shared-memory load overhead is unchanged. The separable kernel is slowest "
    "here because two kernel launches carry a fixed dispatch overhead that dominates at this "
    "arithmetic intensity. Effective bandwidth for the naive kernel reaches 374 GB/s on 2048\u00d72048, "
    "corresponding to 146% of the 256 GB/s peak\u2014consistent with L2 cache hits contributing "
    "to the apparent bandwidth."
)

subsec_h('6.6.2', 'Medium Kernels (r = 3, r = 7)')
para(
    "From r = 3 onwards the tiled kernel consistently outperforms naive. At r = 7 on the "
    "largest image, tiled GPU is 1.61\u00d7 faster than naive and the separable kernel is "
    "2676\u00d7 faster than the CPU reference. The growing advantage of separable is explained "
    "by its O(2K) arithmetic per pixel versus O(K\u00b2) for the 2-D kernels: at K = 15 this "
    "is a 7.5\u00d7 arithmetic saving in addition to the bandwidth saving from tiling."
)

subsec_h('6.6.3', 'Large Kernels (r = 15)')
para(
    "At r = 15 the separable kernel dominates by a large margin. Its execution time on a "
    "2048\u00d72048 image is 0.365 ms, compared to 3.682 ms for tiled and 6.095 ms for naive. "
    "The tiled 2-D kernel at this radius has an arithmetic intensity of 240 FLOP/byte, "
    "placing it well above the device ridge point of ~63 FLOP/byte (16.1 TFLOP/s \u00f7 256 GB/s), "
    "so it is compute-bound rather than bandwidth-bound. The separable kernel, at 15.5 FLOP/byte, "
    "remains bandwidth-bound but issues far fewer floating-point operations, finishing in "
    "a fraction of the time. The highest recorded speedup over the CPU reference was 8495\u00d7 "
    "for the separable kernel at r = 15 on the 2048\u00d72048 image."
)

subsec_h('6.6.4', 'Correctness')
para(
    "Every GPU result passed the per-pixel error check (max absolute difference < 10\u207b\u00b3). "
    "The measured errors ranged from 3.6\u00d710\u207b\u2077 for the naive and tiled kernels to "
    "3.3\u00d710\u207b\u2076 for the separable kernel at r = 15. The slightly larger separable "
    "error arises from floating-point rounding in two sequential passes rather than one, and "
    "is well within acceptable tolerance."
)

sec_h('6.7', 'Roofline Analysis')
para(
    "Table 6.2 summarises the arithmetic intensity and the observed performance regime for each "
    "method at each radius. The device ridge point is approximately 63 FLOP/byte "
    "(16.1 TFLOP/s \u00f7 256 GB/s for the RTX 4070 Laptop GPU)."
)
tbl(['Method', 'r', 'AI (FLOP/B)', 'Regime', 'Observation'],
    [
        ('Naive / Tiled 2-D', 1,   '2.2',   'Memory BW',  'Both BW-limited; naive is faster due to lower overhead'),
        ('Separable',         1,   '1.5',   'Memory BW',  'Launch overhead dominates at this tiny kernel'),
        ('Naive / Tiled 2-D', 3,   '12.2',  'Memory BW',  'Tiled begins to show advantage'),
        ('Separable',         3,   '3.5',   'Memory BW',  'Comparable to tiled on 512\u00d7512; faster at 1024+'),
        ('Naive / Tiled 2-D', 7,   '56.2',  'Compute',    'Just above ridge; tiled 1.6\u00d7 over naive'),
        ('Separable',         7,   '7.5',   'Memory BW',  'BW-limited; fastest among GPU methods'),
        ('Naive / Tiled 2-D', 15,  '240.2', 'Compute',    'Deeply compute-bound; tiled 1.66\u00d7 over naive'),
        ('Separable',         15,  '15.5',  'Memory BW',  'BW-limited; 10\u00d7 faster than tiled at 2048\u00d72048'),
    ],
    col_w=[1.5, 0.4, 0.9, 1.0, 2.6])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  CHAPTER 7
# ══════════════════════════════════════════════════════════════════════════════
ch(7, 'Contribution Summary')
para(
    "Work was divided among the three group members as outlined in Table 7.1. "
    "Design decisions and final testing were carried out jointly."
)
tbl(['Member', 'Reg. No.', 'Responsibility'],
    [
        ('Vasishta Nandipati', '230962095',
         'Naive CUDA kernel (naive_conv.cu), benchmark driver and CPU reference (main.cu), '
         'hardware setup and build configuration'),
        ('Rishit Mathur', '220962091',
         'Shared-memory tiled kernel (tiled_conv.cu), halo and boundary handling, '
         'correctness verification, Nsight Compute profiling'),
        ('Krishiv Kolanu', '230962023',
         'Separable kernel pipeline (separable_conv.cu), Gaussian kernel generation, '
         'roofline analysis, final report'),
    ],
    col_w=[1.7, 1.0, 3.7])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  REFERENCES
# ══════════════════════════════════════════════════════════════════════════════
para('REFERENCES', WD_ALIGN_PARAGRAPH.CENTER, 15, bold=True, sb=12, sa=12)
refs = [
    "V. Podlozhnyuk, \u201cImage Convolution with CUDA,\u201d NVIDIA GPU Computing SDK White Paper, NVIDIA Corporation, Santa Clara, CA, USA, 2007.",
    "D. B. Kirk and W. W. Hwu, Programming Massively Parallel Processors: A Hands-on Approach, 4th ed. Morgan Kaufmann / Elsevier, Waltham, MA, USA, 2022.",
    "P. Getreuer, \u201cA Survey of Gaussian Convolution Algorithms,\u201d Image Processing On Line, vol. 3, pp. 286\u2013310, Nov. 2013. doi: 10.5201/ipol.2013.87.",
    "S. Hong, H. Kim, and J. Lee, \u201cWarp-Level Primitives for Efficient Reduction and Convolution on NVIDIA Ampere GPUs,\u201d in Proc. IEEE IPDPS, Lyon, France, May 2023, pp. 614\u2013624.",
    "N. Vasilache et al., \u201cTensor Comprehensions: Framework-Agnostic High-Performance Machine Learning Abstractions,\u201d arXiv preprint arXiv:1802.04730, Feb. 2018.",
    "S. Chetlur et al., \u201ccuDNN: Efficient Primitives for Deep Learning,\u201d arXiv preprint arXiv:1410.0759, Oct. 2014.",
    "NVIDIA Corporation, CUDA C++ Programming Guide, Version 12.4, Santa Clara, CA, USA, 2024.",
    "W. Luk and T. Voss, \u201cOptimising Convolution Operations on FPGAs and GPUs Using Roofline-Guided Design,\u201d in Proc. IEEE FPT, Tianjin, China, Dec. 2022, pp. 1\u20138.",
    "S. Williams, A. Waterman, and D. Patterson, \u201cRoofline: An Insightful Visual Performance Model for Multicore Architectures,\u201d Commun. ACM, vol. 52, no. 4, pp. 65\u201376, Apr. 2009.",
    "G. Bradski, \u201cThe OpenCV Library,\u201d Dr. Dobb\u2019s Journal of Software Tools, vol. 25, pp. 120\u2013125, Nov. 2000.",
]
for i, ref in enumerate(refs, 1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.35); p.paragraph_format.first_line_indent = Inches(-0.35)
    sfont(p.add_run(f"[{i}]  {ref}"), 12)

# ── Save ──────────────────────────────────────────────────────────────────────
out = '/home/pixie/Documents/sem6/pcap/mini-proj/report/Final_Report.docx'
doc.save(out)
print(f"Saved: {out}")
